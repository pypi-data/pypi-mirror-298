import os
import json
import shutil
from docx import Document
from ..abstract_handler import AbstractHandler

class MicrosoftWordReaderHandler(AbstractHandler):
    
    def handle(self, request: dict) -> dict:
        print("Processing DOCX file...")

        # Initialize variables
        text_content = ''
        media_files = []

        # Load the document from the path specified in the request
        document_path = request.get("path", None)
        document = Document(document_path)

        # Determine the write file path and output folder
        write_file_path = request.get("write_file_path", None)
        if write_file_path:
            base_folder = os.path.dirname(write_file_path)
            file_id = os.path.basename(write_file_path).split('.')[0]
            output_folder = os.path.join(base_folder, file_id)
            media_folder = os.path.join(output_folder, "media")
            os.makedirs(media_folder, exist_ok=True)
        
        # Initialize the metadata dictionary
        if "metadata" not in request:
            request["metadata"] = {}

        # Function to extract and save images, audio, and video
        def extract_and_save_media(run, paragraph_id, source_type, **kwargs):
            # Extract images
            if run.element.xpath('.//a:blip'):
                for blip in run.element.xpath('.//a:blip'):
                    rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    image_part = document.part.related_parts[rId]
                    image_ext = image_part.content_type.split('/')[1]
                    image_file_name = f"{source_type}_image_{rId}.{image_ext}"
                    image_file_path = os.path.join(media_folder, image_file_name)
                    with open(image_file_path, 'wb') as f:
                        f.write(image_part.blob)
                    media_files.append({"type": "image", "path": image_file_path, "paragraph_id": paragraph_id, **kwargs})
            # Extract videos and audio
            if run.element.xpath('.//a:videoFile'):
                for video in run.element.xpath('.//a:videoFile'):
                    rId = video.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}link')
                    video_part = document.part.related_parts[rId]
                    video_ext = video_part.content_type.split('/')[1]
                    video_file_name = f"{source_type}_video_{rId}.{video_ext}"
                    video_file_path = os.path.join(media_folder, video_file_name)
                    with open(video_file_path, 'wb') as f:
                        f.write(video_part.blob)
                    media_files.append({"type": "video", "path": video_file_path, "paragraph_id": paragraph_id, **kwargs})

            if run.element.xpath('.//a:audioFile'):
                for audio in run.element.xpath('.//a:audioFile'):
                    rId = audio.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}link')
                    audio_part = document.part.related_parts[rId]
                    audio_ext = audio_part.content_type.split('/')[1]
                    audio_file_name = f"{source_type}_audio_{rId}.{audio_ext}"
                    audio_file_path = os.path.join(media_folder, audio_file_name)
                    with open(audio_file_path, 'wb') as f:
                        f.write(audio_part.blob)
                    media_files.append({"type": "audio", "path": audio_file_path, "paragraph_id": paragraph_id, **kwargs})

        extract_media = request.get("extract_media", False)

        # Iterate through each paragraph and run in the document
        for paragraph_id, paragraph in enumerate(document.paragraphs):
            text_content += paragraph.text + '\n'
            if extract_media:
                for run_id, run in enumerate(paragraph.runs):
                    extract_and_save_media(run, paragraph_id, source_type="paragraph", run_id=run_id)

        # Iterate through each table in the document
        for table_id, table in enumerate(document.tables):
            for row_id, row in enumerate(table.rows):
                for cell_id, cell in enumerate(row.cells):
                    for paragraph_id, paragraph in enumerate(cell.paragraphs):
                        text_content += paragraph.text + '\n'
                        if extract_media:
                            for run_id, run in enumerate(paragraph.runs):
                                extract_and_save_media(run, paragraph_id, source_type="table", table_id=table_id, row_id=row_id, cell_id=cell_id, run_id=run_id)

        # Save the transcript to a text file
        transcript_file_path = os.path.join(output_folder, 'transcript.txt')
        with open(transcript_file_path, 'w') as transcript_file:
            transcript_file.write(text_content)

        # Create metadata with traceability for graph database
        metadata = {
            "original_file": document_path,
            "transcript_file": transcript_file_path,
            "media_files": media_files
        }
        
        # Save metadata to a json file
        metadata_file_path = os.path.join(output_folder, 'metadata.json')
        with open(metadata_file_path, 'w') as metadata_file:
            json.dump(metadata, metadata_file, indent=4)

        # Update the request metadata with the new metadata
        request["metadata"][document_path] = metadata

        # Update the request with the extracted text
        request.update({"text": text_content})

        # Call the next handler in the chain
        return super().handle(request)
