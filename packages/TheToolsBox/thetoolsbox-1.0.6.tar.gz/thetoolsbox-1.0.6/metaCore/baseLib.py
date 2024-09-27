from docx import Document
import openpyxl
import os
import psutil


def create_excel(data, filename="components.xlsx"):
    # Crée un classeur
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Composants et Prix"

    # Ajoute des en-têtes
    sheet.append(["Composant", "Prix (€)"])

    # Ajoute des données
    for item, price in data:
        sheet.append([item, price])

    # Sauvegarde le fichier
    workbook.save(filename)
    print(f"Fichier Excel '{filename}' créé avec succès.")


def create_word(title, content, filename="rapport_projet.docx"):
    doc = Document()

    # Ajouter un titre
    doc.add_heading(title, 0)

    # Ajouter du contenu
    doc.add_paragraph(content)

    # Sauvegarde le fichier Word
    doc.save(filename)
    print(f"Fichier Word '{filename}' créé avec succès.")


def clean_duplicates(liste):
    outList = []
    for i in liste:
        if i not in outList:
            outList.append(i)
    return outList


def testOpen(root):
    out = False
    try:
        out = os.listdir(root)
        if out:
            out = True
    except:
        # print (f"{root} can't be opened")
        pass

    return out


def disk():
    part = []
    for disk in psutil.disk_partitions():
        part.append(disk[0])
    return part


def typeFile(fileOrDos):
    td = None
    if len(fileOrDos.split('.')) == 1:
        td = '.dos'
    elif len(fileOrDos.split('.')[-1]) > 4:
        td = '.dos'
    else:
        td = '.' + fileOrDos.split('.')[-1]
    return td


def SearchIn (path, target, cycle = 0, listPath = None, dos = 0, file = 0):
    print (path)

    def testOpen(root):
        out = False
        try:
            out = os.listdir(root)
            if out:
                out = True
        except:
            # print (f"{root} can't be opened")
            pass

        return out


def ListingDosFile(path, target=None, cycle=0, listPathDos=None, listPathFile=None):

    if cycle == 0:
        listPathDos = []
        listPathFile = []
    cycle += 1
    if target:
        if target in path:
            if typeFile(path.split('\\')[-1]) == '.dos':
                listPathDos.append(path)
            else:
                listPathFile.append(path)
    else:
        if typeFile(path.split('\\')[-1]) == '.dos':
            listPathDos.append(path)
        else:
            listPathFile.append(path)

    if testOpen(path) and typeFile(path.split('\\')[-1]) == '.dos':
        for dosIn in os.listdir(path):
            if cycle > 0 and path[-1] != '\\':
                path += '\\'
            ListingDosFile(path + dosIn, target, cycle, listPathDos, listPathFile)

    return listPathDos, listPathFile, len(listPathDos), len(listPathFile)
