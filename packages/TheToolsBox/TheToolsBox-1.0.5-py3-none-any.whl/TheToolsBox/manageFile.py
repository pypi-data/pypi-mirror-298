
from docx import Document
import openpyxl

def create_excel(data, filename="components.xlsx"):
    # Crée un classeur
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Composants et Prix"

    # Ajoute des en-têtes
    sheet.append(["Composant", "Prix (€)"])

    # Ajoute des données
    for item, price in data:
        sheet.append([item, price])

    # Sauvegarde le fichier
    workbook.save(filename)
    print(f"Fichier Excel '{filename}' créé avec succès.")
def create_word(title, content, filename="rapport_projet.docx"):
    doc = Document()

    # Ajouter un titre
    doc.add_heading(title, 0)

    # Ajouter du contenu
    doc.add_paragraph(content)

    # Sauvegarde le fichier Word
    doc.save(filename)
    print(f"Fichier Word '{filename}' créé avec succès.")


