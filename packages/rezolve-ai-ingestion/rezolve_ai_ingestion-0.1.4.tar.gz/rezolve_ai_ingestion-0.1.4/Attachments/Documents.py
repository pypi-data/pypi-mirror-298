import os
import tempfile
import textract
from docx import Document
from io import BytesIO

class DocumentReader:
    @staticmethod
    def read_docx(file_content: BytesIO) -> str:
        try:
            doc = Document(file_content)
            return '\n'.join(para.text for para in doc.paragraphs)
        except Exception as e:
            print(f"Error reading .docx: {e}")
            return None

    @staticmethod
    def convert_to_docx(file_content: BytesIO, file_name: str) -> BytesIO:
        # Convert non-docx file to docx using textract
        try:
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file_name)[1])
            temp_file.write(file_content.getvalue())
            temp_file.close()

            # Extract text using textract
            text = textract.process(temp_file.name).decode('utf-8')

            # Create a docx file from the extracted text
            docx_io = BytesIO()
            doc = Document()
            doc.add_paragraph(text)
            doc.save(docx_io)
            docx_io.seek(0)
            return docx_io
        except Exception as e:
            print(f"Error converting to .docx: {e}")
            return None
        finally:
            os.unlink(temp_file.name)

    @staticmethod
    def process_file_content(file_content: BytesIO, file_name: str) -> str:
        extension = file_name.lower().split('.')[-1]
        if extension == 'docx':
            return DocumentReader.read_docx(file_content)
        else:
            docx_content = DocumentReader.convert_to_docx(file_content, file_name)
            return DocumentReader.read_docx(docx_content)
