import os
import re
import csv
import xlrd
from PIL import Image
import pandas as pd



# **************************
# 表格相关
# **************************

# AI优化
def diclis_to_csv(dic_list,csv_file):
    '''给出python代码，输入字典元素的列表，以键为表头，值为数据，输出csv'''
    import csv
    with open(csv_file, 'w', encoding='utf-8') as f:
        # 获取字典元素的键列表
        fieldnames = dic_list[0].keys()
        # 创建csv写入对象
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        # 写入表头
        writer.writeheader()
        # 写入数据
        writer.writerows(dic_list)

# AI优化 #rowlis为列数组
def excel_to_dic(srcx,rdic=True,tablenum=0,regexStr=None,rowlis=None):
    '''
    该函数用于将excel表格转换为字典或列表，
    其中参数srcx为excel表格文件路径，
    rdic为布尔值，若为True，则返回字典，若为False，则返回列表；
    tablenum为表格编号；
    regexStr为正则表达式，用于过滤某些行；
    rowlis为一个列表，用于指定返回的列表中的列索引
    '''
    xlsx = xlrd.open_workbook(srcx,on_demand=True)
    sheet1 = xlsx.sheets()[tablenum]
    lis=[]
    dic={}
    inx=list(str(x.value) for x in sheet1.row(0))
    for i in range(len(inx.copy())):
        if inx[:i+1].count(inx[i])>1:
            inx[i]+='.'+str(inx[:i+1].count(inx[i]))
    for i in range(0,sheet1.nrows):
        if regexStr is None or re.match(regexStr, sheet1.row(i)[0].value) or i==0:
            tmp=list(x.value for x in sheet1.row(i))
            if i!=0 and tmp[0]!='':
                dic[tmp[0]]={}
                for y in range(len(tmp[1:])):
                    dic[tmp[0]][inx[y]]=tmp[y]
            if rowlis:
                tmp=list(tmp[ix] for ix in rowlis)
            lis.append(tmp)
    if rdic:
        return dic
    return lis


# AI优化
def xlsx_to_table(xlsfile,isString = False):
    '''给出python代码，转换xlsx的第一张表格到utf-8编码的csv'''
    workbook = xlrd.open_workbook(xlsfile)
    sheet = workbook.sheet_by_index(0)
    nl=[]
    for row_num in range(sheet.nrows):
        row_value = sheet.row_values(row_num)
        # 可选
        if isString == True:
            for strd in row_value:
                row_value[row_value.index(strd)] = str(strd).strip()
        nl.append(list(row_value))
    return nl



def xlsx_to_list(src):
    import xlrd
    
    # 打开xlsx文件
    workbook = xlrd.open_workbook(src)
    
    # 获取所有sheet
    sheets = workbook.sheets()
    
    # 获取第一个sheet
    sheet1 = sheets[0]
    
    # 获取第一个sheet的行数和列数
    rows = sheet1.nrows
    cols = sheet1.ncols
    
    # 定义一个二层列表
    list_data = []
    
    # 遍历每一行，把数据添加到二层列表中
    for row in range(rows):
        row_data = []
        for col in range(cols):
            cell_data = sheet1.cell_value(row, col)
            row_data.append(cell_data)
        list_data.append(row_data)
    
    return list_data
    


# AI优化
def xlsx_to_csv(xlsfile,csvfile,isString = False):
    '''给出python代码，转换xlsx的第一张表格到utf-8编码的csv'''
    workbook = xlrd.open_workbook(xlsfile)
    sheet = workbook.sheet_by_index(0)
    with open(csvfile, 'w', encoding='utf-8-sig') as f:
        write = csv.writer(f)
        for row_num in range(sheet.nrows):
            row_value = sheet.row_values(row_num)
            # 可选
            if isString == True:
                for strd in row_value:
                    row_value[row_value.index(strd)] = str(strd).strip()
            write.writerow(row_value)

# AI优化
def csv_to_xlsx(csvfile,xlsxfile):
    '''给出python代码，转换utf-8编码的csv到xlsx'''
    import xlwt
    with open(csvfile, 'r', encoding='utf-8-sig') as f:
        read = csv.reader(f)
        workbook = xlwt.Workbook()
        sheet = workbook.add_sheet('Sheet 1')  # 创建一个sheet表格
        l = 0
        for line in read:
            r = 0
            for i in line:
                sheet.write(l, r, i)  # 一个一个将单元格数据写入
                r = r + 1
            l = l + 1
        workbook.save(xlsxfile)  # 保存Excel

dict_list_to_xls = lambda lis, filename: pd.DataFrame(lis).to_excel(filename, index=False)

dict_list_to_xlsx = dict_list_to_xls

list_to_xlsx = lambda lis,dst : pd.DataFrame(lis[1:], columns=lis[0]).to_excel(dst, index=0)


def excel_set_style(file_name, x, y, Style=None, Font=None):
    '''
    Style like 'Pandas'
    Font : openpyxl.styles.Font
    '''
    import openpyxl
    wb = openpyxl.load_workbook(file_name)
    sheet = wb.active
    if Style:
        sheet.cell(row=x, column=y).style = Style
        """
    设置指定单元格样式
    :param cell: 单元格对象
    :param font_name: 字体名称
    :param font_size: 字体大小
    :param font_bold: 是否加粗
    :param font_italic: 是否斜体
    :param font_underline: 下划线样式
    :param font_color: 字体颜色
    :return:
    """
    if Font:
        sheet.cell(row=x, column=y).font = Font
    wb.save(file_name)





# **************************
# PDF相关
# **************************

def image_to_pdf(image_list, dst): 
    import img2pdf
    open(dst, 'wb').write(img2pdf.convert([Image.open(image) for image in image_list]))

def pdf_to_image(pdf_path, dstdir): 
    import pdf2image
    [image.save(dstdir+'/page_' + str(i) + '.jpg', 'JPEG') for i, image in enumerate(pdf2image.convert_from_path(pdf_path))]

def pdf_rotate_pages(src, output_pdf, rotatelist):
    '''pdf旋转, rotatelist如{0:90,2:180}'''
    from PyPDF2 import PdfWriter, PdfReader
    pdf_writer = PdfWriter()
    pdf_reader = PdfReader(src)
    for i in range(pdf_reader.getNumPages()):
        if i in rotatelist:
            page_i = pdf_reader.getPage(i).rotateClockwise(rotatelist[i])
            pdf_writer.addPage(page_i)
        else:
            pdf_writer.addPage(pdf_reader.getPage(i))
    with open(output_pdf, 'wb') as fh:
        pdf_writer.write(fh)

def pdf_split(src):
    '''pdf拆分'''
    from PyPDF2 import PdfWriter, PdfReader
    pdf = PdfReader(src)
    for page in range(pdf.getNumPages()):
        pdf_writer = PdfWriter()
        pdf_writer.add_page(pdf.pages[page])
        output = f'{os.path.splitext(src)[0]}_{str(page).zfill(3)}.pdf'
        with open(output, 'wb') as output_pdf:
            pdf_writer.write(output_pdf)

def pdf_extract(src,s,e):
    '''pdf拆分'''
    from PyPDF2 import PdfWriter, PdfReader
    pdf = PdfReader(src)
    pdf_writer = PdfWriter()
    for page in range(s-1,e):
        pdf_writer.add_page(pdf.pages[page])
    output = f'{os.path.splitext(src)[0]}_{s}_{e}.pdf'
    with open(output, 'wb') as output_pdf:
        pdf_writer.write(output_pdf)


def pdf_merge(srclist, output):
    '''pdf合并'''
    from PyPDF2 import PdfWriter, PdfReader
    pdf_writer = PdfWriter()
    for src in srclist:
        pdf_reader = PdfReader(src)
        for page in range(len(pdf_reader.pages)):
        # 将每页添加到writer对象
            pdf_writer.add_page(pdf_reader.pages[page])
    # 写入合并的pdf
    with open(output, 'wb') as out:
        pdf_writer.write(out)

def pdf_encrypt(input_pdf, output_pdf, password):
    '''pdf加密'''
    from PyPDF2 import PdfWriter, PdfReader
    pdf_writer = PdfWriter()
    pdf_reader = PdfReader(input_pdf)
    for page in range(pdf_reader.getNumPages()):
        pdf_writer.addPage(pdf_reader.getPage(page))
    pdf_writer.encrypt(user_pwd=password, owner_pwd=None,use_128bit=True)
    with open(output_pdf, 'wb') as fh:
        pdf_writer.write(fh)

def pdf_decrypt(input_pdf, output_pdf, password):
    '''pdf解密'''
    from PyPDF2 import PdfWriter, PdfReader
    pdf_writer = PdfWriter()
    pdf_reader = PdfReader(input_pdf)
    pdf_reader.decrypt(password=password)
    for page in range(pdf_reader.getNumPages()):
        pdf_writer.addPage(pdf_reader.getPage(page))
    with open(output_pdf, 'wb') as fh:
        pdf_writer.write(fh)

def pdf_create_watermark(input_pdf, output, watermark):
    '''pdf水印，watermark也是pdf'''
    from PyPDF2 import PdfWriter, PdfReader
    watermark_obj = PdfReader(watermark)
    watermark_page = watermark_obj.getPage(0)
    pdf_reader = PdfReader(input_pdf)
    pdf_writer = PdfWriter()
    # 给所有页面添加水印
    for page in range(pdf_reader.getNumPages()):
        page = pdf_reader.getPage(page)
        page.mergePage(watermark_page)
        pdf_writer.addPage(page)
    with open(output, 'wb') as out:
        pdf_writer.write(out)

def pdf_search(src,word):
    import PyPDF2
    pdfFileObj = open(src, 'rb')
    pdfReader = PyPDF2.PdfReader(pdfFileObj)
    for i in range(0, pdfReader.numPages):
        pageObj = pdfReader.getPage(i)
        if word.upper() in pageObj.extractText().upper():
            return True
    return False

def html_to_pdf(src,dst):
    import pdfkit
    # 将html转换为pdf
    pdfkit.from_file(src, dst)

def url_to_pdf(url,dst):
    import pdfkit
    pdfkit.from_url(url, dst)


def pdf_to_word(in_file,out_file):
    import comtypes.client

    # 创建word应用
    word = comtypes.client.CreateObject('Word.Application')
    
    # 打开pdf文件
    doc = word.Documents.Open(in_file)
    
    # 保存word文件
    doc.SaveAs(out_file, FileFormat=16)
    
    # 关闭word文件
    doc.Close()


# **************************
# Markdown相关
# **************************

def md_to_html(src,dst):
    import markdown
    
    # 读取markdown文件
    with open(src, 'r', encoding='u8') as f:
        content = f.read()
    
    # 转换markdown文件到html文件
    html = markdown.markdown(content)
    
    # 写入html文件
    with open(dst, 'w', encoding='u8') as f:
        f.write(html)


def html_to_markdown(src,dst):
    from markdownify import markdownify
    
    with open(src, 'r', encoding='u8') as f:
        html_data = f.read()
    
    markdown_data = markdownify(html_data)
    
    with open(dst, 'w', encoding='u8') as f:
        f.write(markdown_data)


def md_to_docx_pandoc(md_file):
    '''
    该函数用于将markdown文件转换成docx文件
    参数：
        md_file：markdown文件路径
    返回值：
        无
    '''
    # 导入需要的模块
    import pypandoc
    import os
    
    # 获取文件名
    file_name = os.path.basename(md_file)
    # 将文件名中的.md替换成.docx
    docx_name = file_name.replace(".md", ".docx")
    # 调用pypandoc模块将markdown文件转换为docx文件
    pypandoc.convert_file(md_file, 'docx', outputfile=docx_name)

def md_to_docx_markdown2(md_file, docx_file):
    import markdown2
    import docx
    # Read Markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        md_text = f.read()
    # Convert Markdown to HTML
    html_text = markdown2.markdown(md_text)
    # Create a new document
    document = docx.Document()
    # Add HTML content to the document
    document.add_paragraph(html_text)
    # Save the document
    document.save(docx_file)


def docx_to_markdown(docx_file):
    """
    Convert a docx file to markdown format

    Parameters
    ----------
    docx_file : str
        The path to the docx file

    Returns
    -------
    str
        The markdown version of the docx file
    """
    import docx
    from markdownify import markdownify

    doc = docx.Document(docx_file)
    markdown_text = ''

    for para in doc.paragraphs:
        markdown_text += markdownify(para.text) + '\n'

    return markdown_text



# **************************
# Word相关
# **************************




# **************************
# SQLITE相关
# **************************

def query_sqlite_db(db_file, sql_cmd):
    import sqlite3
    conn = sqlite3.connect(db_file)
    cursor = conn.cursor()
    cursor.execute(sql_cmd)
    records = cursor.fetchall()
    conn.close()
    return records

def sqlite_to_csv(db_file,dstdir):
    import sqlite3
    import pandas as pd
    
    # 连接数据库
    conn = sqlite3.connect(db_file)
    
    # 获取数据库中所有表格
    cursor = conn.execute("SELECT name FROM sqlite_master WHERE type='table';")
    tables = [v[0] for v in cursor.fetchall()]
    
    # 遍历表格，将每一张表格保存为csv和excel文件
    for t in tables:
        df = pd.read_sql_query("SELECT * from %s" % t, conn)
        df.to_csv(dstdir+'/%s.csv' % t, index=False)
        df.to_excel(dstdir+'/%s.xlsx' % t, index=False)
    
    # 关闭数据库连接
    conn.close()

